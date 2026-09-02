import os
import uuid
from typing import List, Dict, Any
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from pydantic import BaseModel
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document

router = APIRouter(prefix="/knowledge", tags=["Knowledge Base"])

def get_vectorstore():
    db_dir = os.path.join(os.path.dirname(__file__), "..", "chroma_db")
    embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-2")
    return Chroma(
        persist_directory=db_dir, 
        embedding_function=embeddings,
        collection_name="dlf_properties"
    )

class TextFactInput(BaseModel):
    facts: List[str]
    source: str

@router.get("/")
async def list_facts():
    try:
        vs = get_vectorstore()
        # Chroma doesn't have a simple "get all" if the DB is large, 
        # but for a PoC we can get the collection and fetch all items
        collection = vs._collection
        results = collection.get()
        
        facts = []
        for i in range(len(results['ids'])):
            facts.append({
                "id": results['ids'][i],
                "content": results['documents'][i],
                "source": results['metadatas'][i].get("source", "Unknown") if results['metadatas'][i] else "Unknown"
            })
            
        return {"facts": facts}
    except Exception as e:
        return {"error": str(e), "facts": []}

@router.post("/")
async def add_text_facts(input_data: TextFactInput):
    try:
        vs = get_vectorstore()
        
        documents = [
            Document(page_content=fact.strip(), metadata={"source": input_data.source}) 
            for fact in input_data.facts if fact.strip()
        ]
        
        if documents:
            vs.add_documents(documents)
            
        return {"status": "success", "added": len(documents)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/{fact_id}")
async def delete_fact(fact_id: str):
    try:
        vs = get_vectorstore()
        vs.delete([fact_id])
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/stats")
async def get_stats():
    try:
        vs = get_vectorstore()
        collection = vs._collection
        results = collection.get()
        
        total_facts = len(results['ids'])
        sources = list(set([
            m.get("source", "Unknown") for m in results['metadatas'] if m
        ])) if results['metadatas'] else []
        
        return {
            "total_facts": total_facts,
            "sources": sources,
            "source_count": len(sources)
        }
    except Exception as e:
        return {"total_facts": 0, "sources": [], "source_count": 0, "error": str(e)}

@router.post("/upload")
async def upload_document(file: UploadFile = File(...), source: str = Form(...)):
    try:
        import PyPDF2
        import docx
        import tempfile
        
        # Save uploaded file temporarily
        suffix = os.path.splitext(file.filename)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name
            
        extracted_text = ""
        
        if suffix == ".pdf":
            with open(temp_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n\n"
        elif suffix == ".docx":
            doc = docx.Document(temp_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        else:
            os.remove(temp_path)
            raise HTTPException(status_code=400, detail="Unsupported file format. Only PDF and DOCX are allowed.")
            
        os.remove(temp_path)
        
        # Simple chunking (split by double newlines or chunks of ~200 words)
        paragraphs = [p.strip() for p in extracted_text.split("\n\n") if len(p.strip()) > 30]
        chunks = []
        for p in paragraphs:
            # Further split if too long (rough word count)
            words = p.split()
            if len(words) > 200:
                for i in range(0, len(words), 150):
                    chunks.append(" ".join(words[i:i+150]))
            else:
                chunks.append(p)
                
        # Embed chunks
        vs = get_vectorstore()
        documents = [
            Document(page_content=chunk, metadata={"source": source or file.filename}) 
            for chunk in chunks if chunk.strip()
        ]
        
        if documents:
            vs.add_documents(documents)
            
        return {
            "status": "success", 
            "filename": file.filename,
            "chunks_extracted": len(documents)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
