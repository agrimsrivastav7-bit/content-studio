import io
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from .. import database as db
import datetime

# PDF generation imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor

# DOCX generation imports
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/export", tags=["Export"])

def get_completed_request(request_id: str):
    req = db.get_request(request_id)
    if not req:
        raise HTTPException(status_code=404, detail="Request not found")
    if req.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Content must be fully approved and validated before export")
    if not req.get("draft_content"):
        raise HTTPException(status_code=400, detail="Draft content is empty")
    return req

@router.get("/{request_id}/pdf")
async def export_pdf(request_id: str):
    req = get_completed_request(request_id)
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'ContentStudioTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        spaceAfter=14,
        textColor=HexColor('#0f172a') # Slate Dark
    )
    meta_style = ParagraphStyle(
        'ContentStudioMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        textColor=HexColor('#64748b'),
        spaceAfter=20
    )
    body_style = ParagraphStyle(
        'ContentStudioBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        leading=16,
        spaceAfter=12
    )
    
    flowables = []
    
    # Header
    flowables.append(Paragraph("CONTENT GOVERNANCE STUDIO", meta_style))
    flowables.append(Paragraph(req.get("topic", "Untitled Document"), title_style))
    
    date_str = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    meta_text = f"Request ID: {request_id} | Generated: {date_str} | Status: Approved"
    flowables.append(Paragraph(meta_text, meta_style))
    flowables.append(Spacer(1, 20))
    
    # Parse markdown-ish draft content into simple paragraphs
    content = req.get("draft_content", "")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        
        # Super simple markdown handling for PDF
        if line.startswith("#"):
            line = line.lstrip("#").strip()
            flowables.append(Paragraph(line, styles['Heading2']))
        else:
            # Handle basic bold
            line = line.replace("**", "<b>").replace("**", "</b>")
            flowables.append(Paragraph(line, body_style))
            
    # Add Risk Score Footer
    flowables.append(Spacer(1, 40))
    scores = req.get("risk_scores", {})
    if scores:
        flowables.append(Paragraph("Validation Summary", styles['Heading3']))
        scores_text = f"Factual: {scores.get('factual_confidence', 'N/A')}% | Compliance: {scores.get('compliance', 'N/A')}% | Brand: {scores.get('brand_alignment', 'N/A')}% | SEO: {scores.get('seo', 'N/A')}%"
        flowables.append(Paragraph(scores_text, meta_style))
    
    doc.build(flowables)
    
    buffer.seek(0)
    
    filename = f"Content_{request_id}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/{request_id}/docx")
async def export_docx(request_id: str):
    req = get_completed_request(request_id)
    
    doc = Document()
    
    # Add Title
    title = doc.add_heading(req.get("topic", "Untitled Document"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add Meta info
    date_str = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph(f"Request ID: {request_id} | Generated: {date_str} | Status: Approved")
    meta.style = doc.styles['Subtitle']
    
    doc.add_paragraph() # spacing
    
    # Add content
    content = req.get("draft_content", "")
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("###"):
            doc.add_heading(line.lstrip("#").strip(), level=3)
        elif line.startswith("##"):
            doc.add_heading(line.lstrip("#").strip(), level=2)
        elif line.startswith("#"):
            doc.add_heading(line.lstrip("#").strip(), level=1)
        else:
            p = doc.add_paragraph(line.replace("**", ""))
            
    # Add footer
    doc.add_page_break()
    doc.add_heading("Validation Summary", level=2)
    scores = req.get("risk_scores", {})
    if scores:
        doc.add_paragraph(f"Factual Confidence: {scores.get('factual_confidence', 'N/A')}%")
        doc.add_paragraph(f"Legal & Compliance: {scores.get('compliance', 'N/A')}%")
        doc.add_paragraph(f"Brand Alignment: {scores.get('brand_alignment', 'N/A')}%")
        doc.add_paragraph(f"SEO Readiness: {scores.get('seo', 'N/A')}%")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"Content_{request_id}.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
